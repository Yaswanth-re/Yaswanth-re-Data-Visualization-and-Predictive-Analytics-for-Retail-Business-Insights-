from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = Path("Retail_Analytics_IEEE_Research_Paper_12_Pages.docx")


def apply_font(run, size=10, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

sect_pr = section._sectPr
cols_list = sect_pr.xpath("./w:cols")
cols = cols_list[0] if cols_list else OxmlElement("w:cols")
if not cols_list:
    sect_pr.append(cols)
cols.set(qn("w:num"), "2")
cols.set(qn("w:space"), "720")

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
styles["Normal"].font.size = Pt(10)
styles["Normal"].paragraph_format.space_after = Pt(3)
styles["Normal"].paragraph_format.line_spacing = 1.0


def add_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    apply_font(r, size=10, bold=True)


def add_para(text, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    apply_font(r, size=10)


def add_ref(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    r = p.add_run(text)
    apply_font(r, size=10)


t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run(
    "Data Visualization and Predictive Analytics for Retail Business Insights: "
    "An Integrated and Interpretable Retail Intelligence Dashboard"
)
apply_font(r, size=16, bold=True)

m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    "Final Year Project Research Paper",
    "Department of Computer Science and Engineering",
    "Prepared in IEEE-style academic format",
]:
    rr = m.add_run(line + "\n")
    apply_font(rr, size=10)

add_heading("Abstract -")
add_para(
    "Retail enterprises operate in an environment shaped by fluctuating demand, seasonal buying behavior, "
    "discount-driven purchasing, regional variation, and changing customer loyalty patterns. These conditions "
    "make it difficult for managers to convert raw transactional data into decisions that are timely, accurate, "
    "and easy to communicate across technical and non-technical teams. This paper presents an integrated retail "
    "analytics framework titled Data Visualization and Predictive Analytics for Retail Business Insights, designed "
    "as an interactive dashboard that unifies descriptive analytics, predictive analytics, customer understanding, "
    "scenario simulation, anomaly-aware monitoring, and automated business insights within a single decision-support "
    "environment. The proposed system accepts retail transaction data in comma-separated value format, performs "
    "structured preprocessing, standardizes key business fields, aggregates records into meaningful analytical views, "
    "and produces interpretable visual outputs for operational review. The framework combines monthly performance "
    "tracking, category and segment comparison, customer segmentation based on behavioral signals, product "
    "recommendation support from co-occurrence patterns, sentiment-oriented review interpretation, and practical "
    "business action summaries. It also compares forecasting models, estimates future performance, and converts "
    "numerical patterns into plain-language findings suitable for managers, students, and presentation audiences. "
    "Particular attention is given to usability, explainability, and graceful handling of missing fields so that "
    "the framework remains functional even when uploaded datasets are incomplete. The final outcome is a "
    "presentation-ready, stakeholder-friendly system that demonstrates how data visualization and predictive "
    "analytics can be combined to improve retail awareness, decision speed, and operational confidence."
)

add_heading("Keywords -")
add_para(
    "retail analytics, business intelligence, data visualization, predictive analytics, sales forecasting, "
    "customer segmentation, scenario simulation, decision support, interactive dashboard, business insights"
)

SECTIONS = [
    (
        "I. INTRODUCTION",
        [
            "Retail businesses generate large volumes of transactional data every day through invoices, online orders, product returns, discounts, customer interactions, and regional store operations. In many organizations, this data remains underused because it is scattered across files and viewed only as isolated reports rather than as a connected business story. Managers often know the sales figure for a given month, yet they may still struggle to answer deeper questions such as why profit is changing, which customer groups are becoming risky, which product areas are driving performance, or what type of planning is needed for the next quarter.",
            "The proposed project addresses this need through a unified dashboard that combines retail data visualization with predictive and explanatory components. Instead of limiting the user to a single chart or isolated metric, the system organizes business understanding into connected modules. These modules allow viewers to inspect performance by time, compare categories and customer segments, study purchasing behavior, inspect discount effects, evaluate product pairings, observe sentiment trends in customer reviews, and examine practical recommendations based on computed patterns.",
            "Another motivation behind the project is the growing demand for decision tools that are both analytical and accessible. Many forecasting or machine learning systems are difficult to explain outside technical teams. In educational and small-business contexts, a dashboard that presents results in simple language can be more valuable than an opaque high-complexity pipeline. This project therefore emphasizes interpretability at every stage through familiar CSV input, simple filters, fallback logic, and readable business output.",
            "The broader significance of this work lies in its integration of descriptive analytics and predictive analytics. Descriptive analytics explains past and current business performance using summaries and visualizations. Predictive analytics extends this by estimating likely future trends and suggesting business actions. By combining both in a single application, the project offers a complete retail intelligence flow that is useful for both business users and academic presentation."
        ],
    ),
    (
        "II. PROBLEM STATEMENT AND RESEARCH MOTIVATION",
        [
            "Retail decision making often suffers from fragmentation. Sales data may exist in spreadsheets, customer comments may be stored in separate systems, and planning may be based on intuition rather than measurable evidence. When leaders do not have a unified view of the business, they can easily overstock slow-moving products, miss high-value customers who are about to disengage, or apply discounts that increase revenue but reduce profit quality.",
            "A second challenge is that many retail datasets are imperfect. Real business files may omit customer identifiers, sub-category labels, review fields, or operational attributes such as quantity and store names. Traditional analytical pipelines often fail when such fields are missing, forcing users to clean data externally before any visualization can begin. Therefore, there is a need for a flexible system that can still deliver useful output under incomplete-data conditions rather than requiring a perfectly prepared dataset.",
            "The third challenge is communication. Even when analytics is technically correct, its business value is limited if the final output is not understandable. Senior managers, academic reviewers, and operational staff usually prefer concise business findings rather than only statistical tables. For that reason, this project treats explanation as part of the system design rather than as an afterthought.",
            "From a research perspective, the project is motivated by the need to evaluate how an integrated dashboard can combine multiple retail analytics tasks into a single workflow. Existing literature often studies forecasting, customer segmentation, anomaly detection, or review analysis separately. In contrast, this project explores how these modules can be orchestrated within one practical interface that supports business monitoring, teaching, and presentation."
        ],
    ),
    (
        "III. OBJECTIVES OF THE STUDY",
        [
            "The first objective of the study is to design a retail dashboard that transforms raw transaction files into meaningful business views. This includes the ability to upload retail CSV files, parse temporal and categorical information, standardize required fields, and present filtered summaries for a chosen business context.",
            "The second objective is to support performance analysis across time, geography, categories, segments, and customers. Retail businesses rarely operate on a single dimension. Decision makers need to know not only total sales, but also which region performs best, which customer group contributes most, which products appear together, and how discount policy shapes profit.",
            "The third objective is to incorporate predictive capability in a form that remains interpretable and practical. Forecasting is valuable only when the resulting numbers help plan inventory, estimate short-term changes, and compare possible business conditions. Accordingly, the system is designed to tie predicted behavior to business actions such as safety stock and scenario comparison.",
            "The fourth objective is to generate plain-language insights for non-technical audiences. This is especially important in educational demonstrations, viva presentations, and managerial reviews. Instead of requiring viewers to infer business meaning from charts, the system identifies notable patterns and explains them in ordinary language."
        ],
    ),
    (
        "IV. REVIEW OF RELATED LITERATURE",
        [
            "Recent literature emphasizes that retail forecasting remains a central task because inventory allocation, procurement timing, staffing, promotions, and revenue planning all depend on the ability to estimate future demand. Taylor and Letham introduced Prophet as a practical forecasting framework for business time series, highlighting the importance of trend decomposition, seasonality, and analyst-friendly configuration in forecasting at scale [1].",
            "Research on retail sales forecasting has also expanded toward machine learning comparisons and model-selection strategies. Studies on retail forecasting using machine learning and optimized hyperparameters report that forecasting performance varies by product category and operational setting, implying that model comparison is often more useful than assuming a universal best approach [2]. Meta-learning work in retail sales forecasting similarly argues that no single forecasting approach fits every SKU or store context [3].",
            "Customer-centric analytics forms another major branch of the literature. Recent work on segmentation using machine learning shows that clustering based on behavioral features such as recency, frequency, and monetary value remains an effective way to divide customers into meaningful strategic groups [4]. Complementary work on RFM-based customer ranking further demonstrates that segmenting customers by observed behavior can support targeted campaigns, retention planning, and improved revenue management [5].",
            "Anomaly detection and external-factor forecasting are increasingly relevant in modern retail due to unstable demand patterns, promotions, and environmental shocks. Liu, Ting, and Zhou introduced Isolation Forest as a scalable method for isolating abnormal observations [6]. Weather-aware retail forecasting studies also show that external signals can materially improve predictive performance and managerial planning [7]."
        ],
    ),
]

REFERENCES = [
    '[1] S. J. Taylor and B. Letham, "Forecasting at Scale," PeerJ Preprints, vol. 5, e3190v2, 2017. [Online]. Available: https://peerj.com/preprints/3190/',
    '[2] M. A. B. et al., "Optimized hyperparameters for retail sales forecasting using grid search," Engineering Applications of Artificial Intelligence, vol. 142, 111472, 2025. [Online]. Available: https://doi.org/10.1016/j.engappai.2025.111472',
    '[3] Y. Ma et al., "Retail sales forecasting with meta-learning," European Journal of Operational Research, vol. 298, no. 1, pp. 24-38, 2022. [Online]. Available: https://www.sciencedirect.com/science/article/pii/S0377221720304847',
    '[4] A. Ashraf, C. A. Rayed, N. A. Awad, and H. M. Sabry, "A Framework for Customer Segmentation to Improve Marketing Strategies Using Machine Learning," Procedia Computer Science, vol. 260, pp. 616-625, 2025. [Online]. Available: https://doi.org/10.1016/j.procs.2025.03.240',
    '[5] A. Joy Christy, A. Umamakeswari, L. Priyatharsini, and A. Neyaa, "RFM ranking - An effective approach to customer segmentation," Journal of King Saud University - Computer and Information Sciences, vol. 33, no. 10, pp. 1251-1257, 2021. [Online]. Available: https://www.sciencedirect.com/science/article/pii/S1319157818304178',
    '[6] F. T. Liu, K. M. Ting, and Z.-H. Zhou, "Isolation forest," in Proc. IEEE International Conference on Data Mining, Pisa, Italy, 2008, pp. 413-422. [Online]. Available: https://research.monash.edu/en/publications/isolation-forest/',
    '[7] H. Chan and M. I. M. Wahab, "A machine learning framework for predicting weather impact on retail sales," Smart Computing and Applications, 100058, 2024. [Online]. Available: https://doi.org/10.1016/j.sca.2024.100058',
    '[8] G. E. P. Box, G. M. Jenkins, G. C. Reinsel, and G. M. Ljung, Time Series Analysis: Forecasting and Control, 5th ed. Hoboken, NJ, USA: Wiley, 2015.',
    '[9] J. Han, M. Kamber, and J. Pei, Data Mining: Concepts and Techniques, 3rd ed. Waltham, MA, USA: Morgan Kaufmann, 2011.',
    '[10] T. Hastie, R. Tibshirani, and J. Friedman, The Elements of Statistical Learning, 2nd ed. New York, NY, USA: Springer, 2009.',
]

SECTIONS.extend(
    [
        (
            "V. SYSTEM OVERVIEW AND PROPOSED FRAMEWORK",
            [
                "The proposed framework is organized as an end-to-end retail analytics pipeline that begins with data intake and ends with business interpretation. The system first loads a CSV file uploaded by the user or, if no file is supplied, falls back to an internal demonstration dataset. Required columns include order date, category, region, sales, and profit. Optional fields such as segment, quantity, discount, customer identifier, review text, and store name are added through safe fallback logic when missing.",
                "After loading, the data passes through a preprocessing stage where the order date is converted into a structured date type, records are sorted chronologically, and user-defined filters are applied. Filters support date range, category, region, segment, target variable, and planning inputs such as scenario assumptions. This filtering step is essential because the value of retail analytics depends strongly on context.",
                "The processed data is then transformed into multiple analytical layers. Time-based aggregation produces monthly series for sales or profit; customer-level grouping produces behavioral features; product co-occurrence analysis produces bundle candidates; review scoring produces sentiment labels; and category or segment summarization produces comparative tables and charts.",
                "Finally, the output layer presents metrics, tables, and narrative findings. Key performance indicators summarize overall business behavior. Charts communicate trends and relationships. Action-focused sections present inventory suggestions and scenario outcomes. The smart insights section converts computed values into sentences such as whether performance is increasing, whether volatility is high, or which region leads the current view."
            ],
        ),
        (
            "VI. DATA PREPROCESSING AND FEATURE PREPARATION",
            [
                "Data preprocessing is one of the most important parts of the project because downstream insights are only as reliable as the input structure. The system begins by validating that the uploaded dataset contains essential business columns required for retail analysis. If mandatory columns such as order date, category, region, sales, or profit are absent, the dashboard stops and informs the user clearly.",
                "For optional fields, the system follows a fault-tolerant strategy. If the segment column is absent, the project inserts a generic segment label so segmentation-based summaries can still function at a basic level. If quantity is missing, the system defaults to a value of one, allowing order-level summaries and operational counts to continue. If customer ID, order ID, store name, sub-category, or review text is missing, synthetic but traceable defaults are generated.",
                "Time preparation is equally important. The order-date field is converted into a proper datetime format and the records are sorted in chronological order. This enables reliable monthly aggregation, growth computation, and time-based comparison. The dashboard groups records at month-end frequency for forecasting-related summaries and for recent trend inspection.",
                "Additional derived features support the customer and product modules. At the customer level, total monetary contribution, purchase frequency, and recency are computed from transactional records. At the product level, grouped order data is used to identify item or sub-category pairs that commonly occur together. Review strings are passed through a simple lexical scoring process to label sentiment as positive, negative, or neutral."
            ],
        ),
        (
            "VII. ANALYTICAL MODULES OF THE PROPOSED SYSTEM",
            [
                "The first module is business overview analytics. This part of the system summarizes monthly behavior, average and peak performance, growth, volatility, and regional contribution. Its purpose is to give users a quick but meaningful starting point. A user can immediately see whether the selected business view is stable, growing, concentrated in a particular region, or supported by specific sub-categories.",
                "The second module is customer intelligence. Customer segmentation groups buyers by behavioral characteristics, while churn-risk labeling identifies customers whose inactivity is materially higher than the typical pattern in the filtered data. Sentiment analysis on available review text adds another layer of interpretation by indicating whether the general tone of customer feedback is positive, neutral, or negative.",
                "The third module covers product, category, and segment analytics. Product recommendation insights identify pairs of products or sub-categories that frequently appear in the same order, which is useful for cross-selling and bundle design. Discount analysis compares average sales and profit across discount ranges, helping users understand whether aggressive discounts truly support profitable growth.",
                "The fourth module is action-oriented analysis and narrative insight generation. Scenario simulation estimates how changes in price, demand, and discount conditions may alter projected revenue. Operational summaries show average order value, quantity movement, and discount intensity. Smart insights convert numeric patterns into plain-language findings that can be used in presentations or executive summaries."
            ],
        ),
        (
            "VIII. FORECASTING, BUSINESS INTERPRETATION, AND SCENARIO SUPPORT",
            [
                "Although the user interface can be tailored to emphasize simpler modules, the project is conceptually grounded in predictive thinking. Monthly retail series are aggregated to create a compact planning view of sales or profit over time. Predictive outputs are used not as isolated numbers but as a basis for business interpretation.",
                "The project's business logic connects predictive estimates to practical actions. A forecast average can be translated into a recommended stock level by applying a safety multiplier that reflects volatility and a chosen stock buffer. If projected change is high or volatility is elevated, the system labels the situation with a stronger reorder-risk warning.",
                "Scenario simulation extends this idea further by allowing decision makers to test hypothetical changes. Instead of waiting to observe a real pricing change or demand shift, the user can adjust planning assumptions and immediately view the effect on revenue-oriented summaries. This is particularly important in retail because pricing, promotions, and expected demand are often under managerial control.",
                "The key contribution here is not merely numeric forecasting, but the translation of predictive output into understandable planning support. This aligns with the broader goal of the project: analytics should guide action. In classroom settings, this demonstrates the practical usefulness of predictive analytics. In business settings, it reduces the gap between model output and decision implementation."
            ],
        ),
    ]
)

SECTIONS.extend(
    [
        (
            "IX. AUTOMATED INSIGHT GENERATION FOR GENERAL AUDIENCES",
            [
                "One of the distinguishing features of the project is the automated insight generator, which converts computed values into plain-language business findings. Traditional dashboards often assume that users can interpret growth values, anomaly counts, volatility levels, or regional rankings on their own. However, many viewers, especially non-technical stakeholders, prefer concise narrative explanations that tell them what the numbers mean.",
                "The importance of this module lies in interpretability and communication. A manager preparing for a review meeting may not want to explain every chart manually. A student presenting a final-year project may want the dashboard to speak clearly to an audience unfamiliar with analytics. In both cases, automated narratives reduce interpretation effort and make the system feel more complete.",
                "Technically, the generator works by combining precomputed analytical indicators with text templates. If the last three monthly values show a monotonic increase or decrease, the system generates a momentum statement. If volatility crosses a threshold, it labels the environment as stable or high-risk. If anomaly periods exist, the system states the count and the latest occurrence.",
                "From a business standpoint, automated insight generation acts as a bridge between analytics and storytelling. Retail decisions are often discussed in meetings, reports, or presentations rather than inside technical notebooks. By producing sentences that summarize important signals, the project makes it easier to communicate findings across departments."
            ],
        ),
        (
            "X. BUSINESS IMPACT AND REAL-WORLD APPLICATIONS",
            [
                "The practical value of the system can be understood through common retail use cases. In inventory planning, a retailer can review monthly behavior, identify strong and weak categories, and use projected averages with safety buffers to estimate stock needs. This helps reduce both overstocking and stockouts.",
                "In marketing and customer relationship management, segmentation and churn-oriented views can guide retention strategy. High-value customers may receive loyalty benefits, early access campaigns, or personalized recommendations. At-risk customers may receive reminder offers or service-focused follow-ups.",
                "Product and discount intelligence also provide meaningful business impact. When the dashboard identifies product pairs that frequently occur together, the retailer can create bundle promotions, shelf placements, or online recommendation rules. When the discount analysis shows that a higher discount band lifts sales but weakens profit, managers can refine pricing policy rather than assuming that more discount always leads to better outcomes.",
                "The project is also valuable in academic and training contexts. It demonstrates a complete analytics workflow using realistic retail concepts and produces outputs that are easy to present in viva examinations, classroom demonstrations, and project exhibitions."
            ],
        ),
        (
            "XI. LIMITATIONS OF THE CURRENT STUDY",
            [
                "Despite its practical strengths, the project has several limitations. First, some modules depend on the richness of the uploaded dataset. If customer IDs, order IDs, discount values, or review text are absent, the system uses fallback values to remain operational. While this protects usability, it also means that some advanced outputs become approximations rather than true reflections of original business records.",
                "Second, the sentiment module uses a simple lexicon-based approach rather than a context-aware natural language model. This makes the implementation lightweight and transparent, but it may miss nuanced statements such as sarcasm, mixed sentiment, or domain-specific wording. Similarly, the automated business insight generator relies on threshold and rule logic.",
                "Third, the forecasting and scenario logic are designed for an educationally accessible dashboard rather than for industrial-scale optimization. Real enterprise forecasting may require hierarchical forecasting, product-level external regressors, event calendars, promotion calendars, and more formal backtesting.",
                "Finally, the system is not currently integrated with live transactional systems, ERP software, or automated scheduling infrastructure. Users must upload data manually or rely on demonstration data. In many real businesses, the next step would be scheduled refresh, role-based access control, audit logging, and persistent storage of historical reports."
            ],
        ),
        (
            "XII. FUTURE IMPROVEMENTS",
            [
                "Future work can extend the system in several meaningful directions. One major improvement would be to include external explanatory variables such as weather, holiday calendars, promotional campaigns, or local events. Prior studies suggest that such features can materially improve retail forecasting accuracy [7].",
                "A second improvement area is model governance and evaluation. Future versions could support rolling-origin backtesting, benchmark comparison against naive forecasts, and automatic model selection rules by category or region. While the current system already compares forecasting performance, a more formal evaluation framework would make the platform stronger for both academic analysis and operational use.",
                "A third area involves richer customer intelligence. Customer lifetime value estimation, next-best-offer recommendations, and campaign-response prediction could be layered on top of the existing segmentation module. The current high-value, regular, and at-risk grouping offers a strong starting point, but future work could integrate loyalty scoring, promotional responsiveness, or return behavior to create more actionable marketing views.",
                "Finally, the user experience can be extended with report scheduling, role-based dashboards for managers and analysts, and automated PDF or presentation exports that include key charts and insight summaries. The project already demonstrates that explainability and accessibility matter."
            ],
        ),
        (
            "XIII. CONCLUSION",
            [
                "This paper presented an integrated framework for Data Visualization and Predictive Analytics for Retail Business Insights, designed as a practical dashboard for turning retail transaction data into understandable and actionable business intelligence. The system combines data preprocessing, descriptive analytics, customer intelligence, product and category comparison, scenario support, and automated insight generation within one presentation-ready environment.",
                "The study demonstrates that retail intelligence is most useful when it is interpretable, robust to imperfect data, and aligned with operational decisions. By allowing users to upload retail files, explore filtered views, inspect customer behavior, compare product and segment performance, and receive narrative findings, the project helps close the gap between raw data and business understanding.",
                "Overall, the project serves as a strong example of how visualization and predictive reasoning can be combined to support better retail decisions. It is suitable as an academic final-year project because it demonstrates data handling, analytics, interface design, business understanding, and research grounding in a single solution.",
            ],
        ),
    ]
)

for heading, paragraphs in SECTIONS:
    add_heading(heading)
    for paragraph in paragraphs:
        add_para(paragraph)

add_heading("REFERENCES")
for reference in REFERENCES:
    add_ref(reference)

add_heading("APPENDIX NOTE")
add_para(
    "This document is written as an original project-specific academic draft prepared for student use. "
    "Exact plagiarism and similarity percentages cannot be guaranteed without running the final file through "
    "the institution's approved plagiarism checker such as Turnitin or iThenticate. However, the paper has "
    "been written in fresh language tailored to the described project and can be further personalized with "
    "student name, college name, dataset details, screenshots, and experimental results before submission."
)

doc.save(OUT)
print(OUT.resolve())
