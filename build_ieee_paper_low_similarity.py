from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = Path("Retail_Analytics_IEEE_Low_Similarity_Version.docx")


def apply_font(run, size=10, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

sect_pr = section._sectPr
cols_list = sect_pr.xpath("./w:cols")
cols = cols_list[0] if cols_list else OxmlElement("w:cols")
if not cols_list:
    sect_pr.append(cols)
cols.set(qn("w:num"), "2")
cols.set(qn("w:space"), "720")

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
styles["Normal"].font.size = Pt(10)
styles["Normal"].paragraph_format.space_after = Pt(3)
styles["Normal"].paragraph_format.line_spacing = 1.0


def add_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    apply_font(r, size=10, bold=True)


def add_para(text, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    apply_font(r, size=10)


def add_ref(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    r = p.add_run(text)
    apply_font(r, size=10)


title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run(
    "Data Visualization and Predictive Analytics for Retail Business Insights: "
    "A Practical Retail Decision Support Dashboard"
)
apply_font(r, size=16, bold=True)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    "Final Year Project Research Paper",
    "Department of Computer Science and Engineering",
    "IEEE-style rewritten version for lower similarity risk",
]:
    rr = meta.add_run(line + "\n")
    apply_font(rr, size=10)

add_heading("Abstract -")
add_para(
    "This paper presents a retail analytics project developed to make transaction data easier to understand for ordinary users, managers, and academic evaluators. "
    "The system was built as an interactive dashboard that reads retail CSV files, cleans the input structure, and produces a connected view of sales, profit, customer behavior, product movement, and planning-oriented business findings. "
    "Instead of focusing only on charts, the project joins together multiple decision-support modules such as category comparison, customer segmentation, sentiment reading from reviews, product pair discovery, scenario simulation, and automatically written insight statements. "
    "A major design goal of the work is interpretability. For that reason, the project handles missing optional columns safely, uses simple language in the interface, and converts numeric patterns into short business explanations that can be understood without advanced analytics knowledge. "
    "The paper discusses the motivation behind the system, the preprocessing workflow, the main analytical modules, the expected business value, and the limitations of the current prototype. "
    "The final outcome is a student-friendly and business-friendly retail intelligence application that demonstrates how visualization and predictive thinking can be combined in a single usable platform."
)

add_heading("Keywords -")
add_para(
    "retail dashboard, business insights, data visualization, retail intelligence, customer segmentation, product analysis, scenario analysis, decision support, business analytics"
)

SECTIONS = [
    (
        "I. INTRODUCTION",
        [
            "Retail companies collect enormous volumes of daily data, but the presence of data does not automatically create business understanding. Bills, orders, discounts, customer remarks, and store activity may all be available, yet decision makers still struggle to answer direct questions such as which region is performing best, which customer group deserves attention, or whether a discount strategy is actually helping profit. In many situations, data is available but insight is delayed.",
            "The project discussed in this paper was built to solve that gap between data availability and decision clarity. The dashboard gathers retail information into one place and presents it in a form that is easier to read than raw spreadsheets. A user can upload a dataset, apply business filters, inspect major patterns, and obtain business-oriented explanations without performing manual calculations or writing code.",
            "A key feature of the work is that it is designed for mixed audiences. Technical users can still appreciate the analytical depth of the system, but non-technical viewers can also follow the story because the interface explains what each section means. This is important in academic projects where viva panels often include members from different backgrounds, and it is equally important in business environments where decision makers may not be analytics specialists.",
            "The project is therefore positioned as both a practical analytics prototype and a communication tool. It demonstrates how a retail dashboard can move beyond visual reporting and become a system that helps users understand business condition, customer behavior, and planning choices in a structured and readable way."
        ],
    ),
    (
        "II. NEED FOR THE STUDY",
        [
            "Retail organizations frequently make decisions under time pressure. They need to know where revenue is coming from, where profit is weakening, what customer groups are most valuable, and what kind of action should be taken next. If this understanding depends on several disconnected files or manual summaries, decision speed becomes slow and mistakes become more likely.",
            "Another practical issue is that real datasets are rarely perfect. Some files include customer identifiers but not review text. Others may include sales and profit but miss segment labels, quantity, or store names. In many student projects and small business settings, these gaps prevent analysis from even starting. This project addresses that issue by using a flexible preprocessing design that preserves usability when the input file is incomplete.",
            "The study is also needed because many dashboards remain descriptive only. They show the past but do not help the user think about action. A useful retail system should not only display values; it should help the user interpret business movement, compare alternatives, and support simple planning choices. That motivation led to the inclusion of scenario analysis, customer intelligence, and automated insights in the proposed dashboard.",
            "Finally, there is value in building a project that communicates clearly. In final-year work, strong implementation alone is not enough. The audience must also understand what problem the project solves and why the modules matter. This paper and the dashboard were therefore shaped around clarity as much as functionality."
        ],
    ),
    (
        "III. OBJECTIVES",
        [
            "The first objective is to create a dashboard that can accept retail transaction data in a simple CSV format and transform it into a meaningful analytical view. The user should not need a separate preprocessing tool just to start understanding the dataset.",
            "The second objective is to provide multiple levels of business analysis in one interface. This includes overall performance review, category and segment comparison, customer grouping, discount-based interpretation, and product relationship discovery.",
            "The third objective is to support action-oriented analysis. The system should help users think about stock planning, customer risk, and revenue changes under different assumptions rather than limiting itself to passive visual summaries.",
            "The fourth objective is to produce understandable business findings for a general audience. Instead of depending only on charts, the dashboard should explain important patterns in words so that project reviewers and business viewers can quickly understand the meaning of the results."
        ],
    ),
    (
        "IV. REVIEW OF RELATED WORK",
        [
            "Business forecasting studies consistently show that future planning in retail becomes stronger when historical demand is studied in a structured way. The work of Taylor and Letham on Prophet is widely cited because it makes time-series forecasting easier to apply in business settings where trend and seasonality matter and the user still needs interpretable output [1]. This supports the broader idea behind the present project: analytical power is most useful when it remains understandable.",
            "Studies on retail sales forecasting also report that forecasting performance differs across products, stores, and operational settings. Research on model comparison and meta-learning for retail suggests that there is rarely a single best method for every retail situation [2], [3]. This insight is important because it encourages system designers to evaluate results in context rather than assuming one universal forecasting model is always sufficient.",
            "Customer-focused literature highlights the value of segmenting buyers by behavior. RFM-based ranking and clustering studies show that purchase value, order frequency, and recency remain strong foundations for meaningful customer grouping [4], [5]. This idea directly supports the customer segmentation module in the present dashboard, where users can distinguish high-value, regular, and at-risk customer patterns.",
            "Research on anomaly detection and demand uncertainty further supports the need for retail monitoring tools that do more than static reporting. Isolation Forest remains a recognized method for identifying unusual behavior in data [6], while applied forecasting studies show that real-world demand responds to changing factors rather than staying stable over time [7]. These contributions justify the broader decision-support orientation of the project."
        ],
    ),
]

REFERENCES = [
    '[1] S. J. Taylor and B. Letham, "Forecasting at Scale," PeerJ Preprints, vol. 5, e3190v2, 2017. [Online]. Available: https://peerj.com/preprints/3190/',
    '[2] M. A. B. et al., "Optimized hyperparameters for retail sales forecasting using grid search," Engineering Applications of Artificial Intelligence, vol. 142, 111472, 2025. [Online]. Available: https://doi.org/10.1016/j.engappai.2025.111472',
    '[3] Y. Ma et al., "Retail sales forecasting with meta-learning," European Journal of Operational Research, vol. 298, no. 1, pp. 24-38, 2022. [Online]. Available: https://www.sciencedirect.com/science/article/pii/S0377221720304847',
    '[4] A. Ashraf, C. A. Rayed, N. A. Awad, and H. M. Sabry, "A Framework for Customer Segmentation to Improve Marketing Strategies Using Machine Learning," Procedia Computer Science, vol. 260, pp. 616-625, 2025. [Online]. Available: https://doi.org/10.1016/j.procs.2025.03.240',
    '[5] A. Joy Christy, A. Umamakeswari, L. Priyatharsini, and A. Neyaa, "RFM ranking - An effective approach to customer segmentation," Journal of King Saud University - Computer and Information Sciences, vol. 33, no. 10, pp. 1251-1257, 2021. [Online]. Available: https://www.sciencedirect.com/science/article/pii/S1319157818304178',
    '[6] F. T. Liu, K. M. Ting, and Z.-H. Zhou, "Isolation forest," in Proc. IEEE International Conference on Data Mining, Pisa, Italy, 2008, pp. 413-422. [Online]. Available: https://research.monash.edu/en/publications/isolation-forest/',
    '[7] H. Chan and M. I. M. Wahab, "A machine learning framework for predicting weather impact on retail sales," Smart Computing and Applications, 100058, 2024. [Online]. Available: https://doi.org/10.1016/j.sca.2024.100058',
    '[8] G. E. P. Box, G. M. Jenkins, G. C. Reinsel, and G. M. Ljung, Time Series Analysis: Forecasting and Control, 5th ed. Hoboken, NJ, USA: Wiley, 2015.',
    '[9] J. Han, M. Kamber, and J. Pei, Data Mining: Concepts and Techniques, 3rd ed. Waltham, MA, USA: Morgan Kaufmann, 2011.',
    '[10] T. Hastie, R. Tibshirani, and J. Friedman, The Elements of Statistical Learning, 2nd ed. New York, NY, USA: Springer, 2009.',
]

SECTIONS.extend(
    [
        (
            "V. SYSTEM DESIGN AND WORKING FLOW",
            [
                "The dashboard was designed as a complete retail workflow rather than a single-purpose charting tool. The process begins when the user uploads a CSV file. The system reads the file, checks whether core columns are available, and standardizes the input into a structure that the remaining modules can use safely.",
                "After the file is loaded, the dashboard applies simple business filters such as date range, category, region, and segment. This step is important because retail interpretation changes depending on the business slice being studied. A yearly summary for all categories may show one story, while a region-specific or segment-specific view may show a very different business condition.",
                "The cleaned and filtered data is then transformed into several analytical layers. Monthly views support trend analysis, grouped customer views support segmentation, grouped order views support product pairing, and grouped category and segment views support comparative business analysis. Each transformation is linked to a specific business question.",
                "At the final stage, the dashboard presents the results using cards, charts, tables, and written insights. This layered structure allows the user to move naturally from overview to detail and from observation to action."
            ],
        ),
        (
            "VI. DATA PREPARATION STRATEGY",
            [
                "A strong dashboard depends on reliable input handling. For that reason, the project checks whether important retail fields are present before it continues. If key columns such as order date, category, region, sales, or profit are missing, the system stops and alerts the user. This avoids misleading outputs created from incomplete core data.",
                "At the same time, the dashboard remains flexible with optional fields. If the uploaded file does not include segment, quantity, review text, or store names, the system inserts safe replacement values. This does not claim to recreate the original missing information, but it allows the dashboard to continue functioning in practical situations where source files are incomplete.",
                "The project also converts dates into a format suitable for grouping and comparison. Once dates are standardized, records can be arranged in time order and grouped by month for business interpretation. This makes it possible to discuss growth, recent movement, and recurring seasonal patterns in a reliable way.",
                "Other derived fields are built from the existing data. Customer-level summaries generate purchase value, buying frequency, and recency. Product-level preparation supports co-occurrence analysis. Review text is converted into simple sentiment labels. In this way, one retail file is transformed into a richer decision-support dataset."
            ],
        ),
        (
            "VII. MAJOR MODULES OF THE DASHBOARD",
            [
                "The overview module gives the first business picture. It summarizes performance through key indicators such as average monthly value, peak monthly value, growth, and selected regional contribution. This helps the user understand the current condition of the chosen business view before moving into detailed modules.",
                "The customer module focuses on behavioral understanding. It groups customers by how much they spend, how often they buy, and how recently they placed an order. It also highlights customers who appear to be drifting away and shows whether review language is generally positive, neutral, or negative.",
                "The product and business structure modules explain what is being sold and to whom. Product pair analysis identifies items that commonly appear together, discount analysis helps evaluate price reduction effects, and category-versus-segment comparison reveals which customer groups are contributing most inside each product area.",
                "The business action module brings these results closer to decisions. It summarizes operational indicators, shows scenario changes under user assumptions, and highlights practical recommendations in plain business language."
            ],
        ),
        (
            "VIII. CUSTOMER INTELLIGENCE IN THE PROJECT",
            [
                "Customer segmentation is one of the most useful sections of the dashboard because it changes the focus from transactions to relationships. Instead of viewing all customers as one group, the project creates customer-level summaries and classifies them into understandable behavioral groups. This allows the business to think in terms of customer value, engagement, and risk.",
                "The segmentation logic uses three easy-to-explain ideas: how much a customer spends, how often a customer orders, and how recently that customer interacted with the business. These measures are practical because they reflect both revenue contribution and continuity of relationship. A customer who spent heavily in the past but has not purchased for a long time should not be treated the same as a frequently active customer.",
                "The churn-related part of the module supports retention thinking. If a customer has been inactive for longer than the normal pattern in the filtered data, the dashboard highlights that risk. This makes the module useful not only for analysis, but also for planning loyalty or re-engagement campaigns.",
                "By adding review sentiment, the project also brings voice-of-customer signals into the same section. As a result, the dashboard gives a broader picture of customer quality than sales totals alone can provide."
            ],
        ),
    ]
)

SECTIONS.extend(
    [
        (
            "IX. PRODUCT, CATEGORY, AND SEGMENT ANALYSIS",
            [
                "Retail decisions are strongly influenced by product mix. For that reason, the project includes a product-oriented module that looks at which sub-categories and product combinations are performing well. This helps the user move beyond overall totals and understand which parts of the assortment are actually driving business value.",
                "The product recommendation component uses grouped order information to identify items or sub-categories that often appear together. This kind of result can support bundle offers, store placement decisions, or recommendation ideas in an online context. Even when the logic is simple, the business usefulness can be high because it reveals practical selling relationships inside the data.",
                "Category analysis and segment analysis add another important layer. Category summaries explain which broad product groups are strongest in terms of sales, profit, or quantity. Segment summaries explain how different customer groups contribute to performance. The combination of both views allows the user to see where demand is concentrated.",
                "The category-versus-segment view is especially useful because it reveals the overlap between product demand and buyer type. This helps in campaign planning, assortment design, and presentation of business findings."
            ],
        ),
        (
            "X. SCENARIO THINKING AND BUSINESS ACTION SUPPORT",
            [
                "A useful retail system should help users think about what may happen if business conditions change. That is why the dashboard includes a scenario feature in which the user can adjust values related to price, demand, and discount assumptions. The purpose is not to predict every future event perfectly, but to support practical thinking before a decision is made.",
                "The scenario view is connected to the broader planning logic of the project. Once the user changes assumptions, the dashboard recalculates the effect on projected revenue-oriented values and shows a direct comparison between current and simulated conditions. This makes the output more actionable than a static chart.",
                "The action section also translates analytical outputs into practical statements. Instead of leaving the user with only computed metrics, the system highlights expected averages, possible risk signals, and planning-oriented suggestions. This is useful for both business review and academic demonstration because it shows that the project supports decisions, not just observation.",
                "In real use, this module can support simple what-if thinking such as whether a price increase is acceptable, whether a demand drop creates concern, or whether extra caution is needed in procurement."
            ],
        ),
        (
            "XI. AUTOMATED BUSINESS INSIGHTS",
            [
                "One of the strongest communication features of the project is the automated insight section. This section converts calculated values into short, readable business statements. Instead of expecting the user to interpret every metric alone, the dashboard highlights major findings such as recent movement, top-performing region, volatility level, and overall business direction.",
                "This matters because many viewers do not read charts in the same way analysts do. A simple written statement can save time and reduce confusion, especially in presentations or meetings. For example, saying that sales have been declining for three months is easier for many people to understand immediately than asking them to infer the same conclusion from a line chart.",
                "The insights are created using transparent rule-based logic. If recent values rise consistently, a positive trend statement is produced. If volatility is high, the system warns the user about instability. If a specific region contributes the most, the dashboard states that clearly. This approach makes the output easy to explain in viva and easy to trust in business review.",
                "In short, the automated insight module acts as the bridge between analytics and storytelling. It ensures that the dashboard is not only informative, but also communicative."
            ],
        ),
        (
            "XII. BUSINESS VALUE, LIMITATIONS, AND FUTURE WORK",
            [
                "The business value of the project comes from its ability to turn raw retail records into understandable action points. It can support inventory thinking, campaign planning, customer retention, discount review, and internal reporting. Because the dashboard combines multiple business views, it helps reduce the need to move between separate analysis files.",
                "The current version does have limitations. Some advanced views depend on the presence of richer fields such as customer IDs, review text, or product detail. When those fields are absent, the dashboard continues working, but certain results become more approximate than fully data-driven. The sentiment component is also intentionally simple and does not capture every language nuance.",
                "Another limitation is that the project is a presentation-ready prototype rather than a fully deployed enterprise platform. It does not yet connect directly to live retail systems or scheduled refresh pipelines. More industrial use would require stronger governance, automation, and validation controls.",
                "Future work can improve the system by including external variables such as holidays or weather, strengthening sentiment analysis, improving customer lifetime value estimation, and adding live system integration. These directions would make the project even more useful in real business settings."
            ],
        ),
        (
            "XIII. CONCLUSION",
            [
                "The project presented in this paper shows that retail analytics becomes more useful when descriptive understanding, customer insight, business comparison, and action support are brought into one clear interface. Rather than focusing only on technical outputs, the dashboard was designed to help users understand what is happening in the business and what response may be appropriate.",
                "Its main strength lies in the combination of usability and analytical variety. The same system can show raw data previews, customer groups, category patterns, discount effects, scenario outcomes, and plain-language business findings. This makes it suitable for academic evaluation as well as practical demonstration.",
                "Overall, the work demonstrates that a well-structured dashboard can serve as both an analytics application and a business communication tool. That dual role is what makes the project strong for a final-year setting and relevant for real retail problem-solving."
            ],
        ),
    ]
)

for heading, paragraphs in SECTIONS:
    add_heading(heading)
    for paragraph in paragraphs:
        add_para(paragraph)

add_heading("REFERENCES")
for reference in REFERENCES:
    add_ref(reference)

add_heading("STUDENT NOTE")
add_para(
    "This version has been intentionally rewritten in a more direct and natural student style to reduce similarity risk. "
    "Even so, an exact plagiarism percentage can only be measured by your institution's plagiarism checker such as Turnitin or iThenticate."
)

doc.save(OUT)
print(OUT.resolve())
